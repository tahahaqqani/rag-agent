import os
import logging
from pathlib import Path
from typing import Iterable, List
from datetime import datetime

from langchain_community.document_loaders import PyPDFLoader, TextLoader, UnstructuredMarkdownLoader
from docx import Document as DocxDocument
from langchain.docstore.document import Document
from .rag import vectorstore

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Chunking configuration
DEFAULT_CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "600"))
DEFAULT_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "80"))


def character_chunks(text: str, max_chars: int = DEFAULT_CHUNK_SIZE, overlap: int = DEFAULT_OVERLAP) -> Iterable[str]:
    """
    Split text into overlapping chunks based on character count.
    
    Args:
        text: Input text to chunk
        max_chars: Maximum characters per chunk
        overlap: Number of overlapping characters between chunks
        
    Yields:
        Text chunks
    """
    try:
        if len(text) <= max_chars:
            yield text
            return
            
        step = max_chars - overlap
        
        for i in range(0, len(text), step):
            chunk_text = text[i:i + max_chars]
            
            # Skip empty chunks
            if chunk_text.strip():
                yield chunk_text
                
    except Exception as e:
        logger.error(f"Error in character chunking: {str(e)}")
        # Fallback to simple chunking
        for i in range(0, len(text), max_chars):
            yield text[i:i + max_chars]


def load_docs(input_dir: str) -> Iterable[Document]:
    """
    Load documents from various file formats.
    
    Args:
        input_dir: Directory containing documents
        
    Yields:
        Document objects with content and metadata
    """
    input_path = Path(input_dir)
    
    if not input_path.exists():
        logger.error(f"Input directory does not exist: {input_dir}")
        return
    
    # Supported file extensions
    supported_extensions = {".pdf", ".txt", ".md", ".markdown"}
    
    for file_path in input_path.rglob("*"):
        if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
            try:
                logger.info(f"Processing file: {file_path}")
                
                if file_path.suffix.lower() == ".pdf":
                    # Handle PDF files
                    loader = PyPDFLoader(str(file_path))
                    pages = loader.load()
                    for i, page in enumerate(pages):
                        # Add page number to metadata
                        page.metadata["page_number"] = i + 1
                        page.metadata["total_pages"] = len(pages)
                        yield page
                        
                elif file_path.suffix.lower() in [".txt", ".md", ".markdown"]:
                    # Handle text and markdown files
                    if file_path.suffix.lower() in [".md", ".markdown"]:
                        loader = UnstructuredMarkdownLoader(str(file_path))
                    else:
                        loader = TextLoader(str(file_path), autodetect_encoding=True)
                    
                    docs = loader.load()
                    for doc in docs:
                        yield doc
                        
                elif file_path.suffix.lower() in [".docx", ".doc"]:
                    # Handle Word documents
                    try:
                        docx = DocxDocument(str(file_path))
                        text = ""
                        for paragraph in docx.paragraphs:
                            if paragraph.text.strip():
                                text += paragraph.text.strip() + "\n\n"
                        
                        # Create a document object
                        yield Document(
                            page_content=text,
                            metadata={"source": str(file_path)}
                        )
                    except Exception as e:
                        logger.error(f"Error processing DOCX file {file_path}: {str(e)}")
                        continue
                        
            except Exception as e:
                logger.error(f"Error processing {file_path}: {str(e)}")
                continue


def ingest_folder(input_dir: str, source_tag: str = "local", chunk_size: int = DEFAULT_CHUNK_SIZE, overlap: int = DEFAULT_OVERLAP) -> int:
    """
    Load documents → chunk → upsert into Chroma.
    
    Args:
        input_dir: Directory containing documents
        source_tag: Source identifier for documents
        chunk_size: Maximum characters per chunk
        overlap: Overlapping characters between chunks
        
    Returns:
        Number of chunks created
    """
    try:
        logger.info(f"Starting ingestion from: {input_dir}")
        logger.info(f"Chunk size: {chunk_size}, Overlap: {overlap}")
        
        docs = []
        total_chunks = 0
        
        for doc in load_docs(input_dir):
            # Create chunks from document content
            for chunk_text in character_chunks(doc.page_content, chunk_size, overlap):
                # Create new document for each chunk
                chunk_doc = Document(
                    page_content=chunk_text,
                    metadata={
                        **(doc.metadata or {}),
                        "source": doc.metadata.get("source", str(doc.metadata.get("source", "")) or source_tag),
                        "chunk_size": chunk_size,
                        "overlap": overlap,
                        "ingested_at": datetime.now().isoformat(),
                        "original_file": doc.metadata.get("source", "unknown"),
                    }
                )
                docs.append(chunk_doc)
                total_chunks += 1
        
        if docs:
            logger.info(f"Adding {len(docs)} chunks to vector store...")
            vectorstore.add_documents(docs)
            # Chroma automatically persists, no need to call persist()
            logger.info(f"Successfully ingested {total_chunks} chunks from {input_dir}")
        else:
            logger.warning(f"No documents found in {input_dir}")
            
        return total_chunks
        
    except Exception as e:
        logger.error(f"Error during ingestion: {str(e)}")
        return 0


def ingest_single_file(file_path: str, source_tag: str = "local", chunk_size: int = DEFAULT_CHUNK_SIZE, overlap: int = DEFAULT_OVERLAP) -> int:
    """
    Ingest a single file into the vector store.
    
    Args:
        file_path: Path to the file
        source_tag: Source identifier
        chunk_size: Maximum characters per chunk
        overlap: Overlapping characters between chunks
        
    Returns:
        Number of chunks created
    """
    try:
        file_path = Path(file_path)
        if not file_path.exists():
            logger.error(f"File does not exist: {file_path}")
            return 0
            
        # Create temporary directory structure
        temp_dir = Path("./temp_ingest")
        temp_dir.mkdir(exist_ok=True)
        
        # Copy file to temp directory
        import shutil
        temp_file = temp_dir / file_path.name
        shutil.copy2(file_path, temp_file)
        
        # Ingest from temp directory
        result = ingest_folder(str(temp_dir), source_tag, chunk_size, overlap)
        
        # Clean up
        shutil.rmtree(temp_dir)
        
        return result
        
    except Exception as e:
        logger.error(f"Error ingesting single file: {str(e)}")
        return 0


def get_ingestion_stats() -> dict:
    """Get statistics about the current vector store."""
    try:
        collection = vectorstore._collection
        if collection:
            count = collection.count()
            return {
                "total_chunks": count,
                "chunk_size": DEFAULT_CHUNK_SIZE,
                "overlap": DEFAULT_OVERLAP,
                "chunking_method": "character-based"
            }
        return {"error": "Collection not initialized"}
    except Exception as e:
        logger.error(f"Error getting ingestion stats: {str(e)}")
        return {"error": str(e)}
